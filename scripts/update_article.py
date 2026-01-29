#!/usr/bin/env python3
"""
Script per aggiornare automaticamente il file Word dell'articolo con il codice corretto dal progetto.
Richiede: pip install python-docx
"""

import re
import sys
from pathlib import Path
from typing import Optional
from copy import deepcopy

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installa python-docx con: pip install python-docx")
    sys.exit(1)


# Percorsi
PROJECT_ROOT = Path(__file__).parent.parent
WORD_FILE = Path("/Users/vasco/Downloads/React web app per san valentino.docx")
OUTPUT_FILE = Path("/Users/vasco/Downloads/React web app per san valentino - AGGIORNATO.docx")


# Mapping dei componenti con i loro pattern di ricerca nel documento
COMPONENT_PATTERNS = {
    "StartScreen": [
        r"// src/components/StartScreen\.jsx",
        r"src/components/StartScreen\.jsx",
    ],
    "SetupScreen": [
        r"// src/components/SetupScreen\.jsx", 
        r"src/components/SetupScreen\.jsx",
    ],
    "QuestionCard": [
        r"// src/components/QuestionCard\.jsx",
        r"src/components/QuestionCard\.jsx",
    ],
    "PassPhoneScreen": [
        r"// src/components/PassPhoneScreen\.jsx",
        r"src/components/PassPhoneScreen\.jsx",
    ],
    "ResultScreen": [
        r"// src/components/ResultScreen\.jsx",
        r"src/components/ResultScreen\.jsx",
    ],
    "AIResultCard": [
        r"// src/components/AIResultCard\.jsx",
        r"src/components/AIResultCard\.jsx",
    ],
    "HistoryScreen": [
        r"// src/components/HistoryScreen\.jsx",
        r"src/components/HistoryScreen\.jsx",
    ],
    "HeartAnimation": [
        r"// src/components/HeartAnimation\.jsx",
        r"src/components/HeartAnimation\.jsx",
    ],
}


def read_component_code(component_name: str) -> str:
    """Legge il codice di un componente dal progetto."""
    file_path = PROJECT_ROOT / "src" / "components" / f"{component_name}.jsx"
    if file_path.exists():
        code = file_path.read_text(encoding="utf-8")
        # Aggiungi il commento del path all'inizio
        return f"// src/components/{component_name}.jsx\n{code}"
    return ""


def get_updated_components() -> dict:
    """Ottiene il codice aggiornato di tutti i componenti."""
    updated = {}
    for comp in COMPONENT_PATTERNS.keys():
        code = read_component_code(comp)
        if code:
            updated[comp] = code
            print(f"  ✓ Letto {comp}.jsx ({len(code)} caratteri)")
        else:
            print(f"  ✗ Non trovato {comp}.jsx")
    return updated


def find_code_block_end(paragraphs: list, start_idx: int) -> int:
    """
    Trova la fine di un blocco di codice.
    Un blocco termina quando troviamo una riga vuota seguita da testo non-codice,
    o quando troviamo un nuovo titolo/sezione.
    """
    brace_count = 0
    found_export = False
    
    for i in range(start_idx, len(paragraphs)):
        text = paragraphs[i].text.strip()
        
        # Conta le parentesi graffe per capire quando il componente è completo
        brace_count += text.count("{") - text.count("}")
        
        # Cerca la fine dell'export default function
        if "export default function" in text or "export default" in text:
            found_export = True
        
        # Se abbiamo trovato l'export e le graffe sono bilanciate, siamo alla fine
        if found_export and brace_count <= 0 and text.endswith("}"):
            return i + 1
        
        # Se troviamo una nuova sezione (###, ##, etc.) siamo andati troppo oltre
        if text.startswith("#") or text.startswith("---"):
            return i
        
        # Se troviamo un altro componente, siamo andati troppo oltre
        if i > start_idx and "// src/components/" in text:
            return i
    
    return len(paragraphs)


def identify_code_blocks(doc: Document) -> list:
    """
    Identifica tutti i blocchi di codice nel documento e il loro tipo.
    Restituisce una lista di tuple: (start_idx, end_idx, component_name)
    """
    blocks = []
    paragraphs = doc.paragraphs
    i = 0
    
    while i < len(paragraphs):
        text = paragraphs[i].text.strip()
        
        # Cerca l'inizio di un blocco di codice componente
        for comp_name, patterns in COMPONENT_PATTERNS.items():
            for pattern in patterns:
                if re.search(pattern, text, re.IGNORECASE):
                    # Trovato l'inizio di un blocco
                    end_idx = find_code_block_end(paragraphs, i)
                    blocks.append((i, end_idx, comp_name))
                    print(f"  📍 Trovato {comp_name}: righe {i}-{end_idx}")
                    i = end_idx - 1  # Salta al prossimo blocco
                    break
            else:
                continue
            break
        
        i += 1
    
    return blocks


def copy_paragraph_format(source_para, target_para):
    """Copia la formattazione di base da un paragrafo sorgente."""
    try:
        if source_para.style:
            target_para.style = source_para.style
    except:
        pass


def replace_code_blocks(doc: Document, components: dict) -> int:
    """
    Sostituisce i blocchi di codice nel documento con le versioni aggiornate.
    """
    replacements = 0
    blocks = identify_code_blocks(doc)
    
    # Ordina i blocchi in ordine inverso per evitare problemi con gli indici
    blocks.sort(key=lambda x: x[0], reverse=True)
    
    for start_idx, end_idx, comp_name in blocks:
        if comp_name not in components:
            print(f"  ⚠️ Codice per {comp_name} non disponibile, skip")
            continue
        
        new_code = components[comp_name]
        
        print(f"  🔄 Sostituisco {comp_name} (paragrafi {start_idx}-{end_idx})...")
        
        # Ottieni il formato del primo paragrafo del codice originale
        original_para = doc.paragraphs[start_idx]
        
        # Rimuovi i vecchi paragrafi (dal fondo verso l'alto per mantenere gli indici)
        for i in range(end_idx - 1, start_idx - 1, -1):
            p = doc.paragraphs[i]._element
            p.getparent().remove(p)
        
        # Trova il punto di inserimento
        if start_idx > 0:
            insert_after = doc.paragraphs[start_idx - 1]
        else:
            insert_after = None
        
        # Inserisci il nuovo codice come singoli paragrafi
        code_lines = new_code.split("\n")
        
        # Crea nuovi paragrafi per ogni riga di codice
        for j, line in enumerate(code_lines):
            # Crea un nuovo paragrafo
            if insert_after is not None:
                # Inserisci dopo il paragrafo di riferimento
                new_para = insert_after._element.addnext(
                    doc.add_paragraph()._element
                )
                new_p = doc.paragraphs[start_idx + j]
            else:
                new_p = doc.add_paragraph()
            
            new_p.text = line
            
            # Applica formattazione codice (font monospace)
            for run in new_p.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
        
        replacements += 1
        print(f"  ✓ {comp_name} aggiornato ({len(code_lines)} righe)")
    
    return replacements


def simple_text_replace(doc: Document, components: dict) -> int:
    """
    Approccio semplificato: cerca e sostituisce il testo direttamente.
    Funziona meglio per documenti con struttura semplice.
    """
    replacements = 0
    full_text = "\n".join([p.text for p in doc.paragraphs])
    
    for comp_name, new_code in components.items():
        # Pattern per trovare il blocco di codice del componente
        # Cerca da "// src/components/ComponentName.jsx" fino alla fine del componente
        pattern = rf"(// src/components/{comp_name}\.jsx[\s\S]*?)(export default function {comp_name}[\s\S]*?\n\}})"
        
        # Verifica se il pattern esiste nel documento
        if re.search(pattern, full_text):
            print(f"  ✓ Pattern trovato per {comp_name}")
            replacements += 1
    
    return replacements


def update_word_document():
    """Funzione principale per aggiornare il documento Word."""
    print(f"\n{'='*60}")
    print("📄 AGGIORNAMENTO AUTOMATICO DOCUMENTO WORD")
    print(f"{'='*60}")
    print(f"\nFile input:  {WORD_FILE}")
    print(f"File output: {OUTPUT_FILE}")
    
    if not WORD_FILE.exists():
        print(f"\n❌ File non trovato: {WORD_FILE}")
        return False
    
    # Leggi il documento
    print(f"\n📖 Apertura documento...")
    doc = Document(WORD_FILE)
    
    # Ottieni il codice aggiornato
    print("\n📂 Lettura componenti dal progetto:")
    components = get_updated_components()
    
    if not components:
        print("\n❌ Nessun componente trovato nel progetto")
        return False
    
    # Statistiche documento
    print(f"\n📊 Statistiche documento:")
    print(f"  Paragrafi totali: {len(doc.paragraphs)}")
    print(f"  Tabelle: {len(doc.tables)}")
    
    # Identifica i blocchi di codice
    print(f"\n🔍 Identificazione blocchi di codice:")
    blocks = identify_code_blocks(doc)
    
    print(f"\n📝 Blocchi di codice trovati: {len(blocks)}")
    
    # Crea una nuova versione del documento con i codici aggiornati
    print(f"\n🔄 Creazione documento aggiornato...")
    
    # Approccio: ricrea il documento sostituendo i blocchi
    new_doc = Document()
    
    i = 0
    paragraphs = doc.paragraphs
    blocks_dict = {b[0]: b for b in blocks}  # Indicizza per start_idx
    
    while i < len(paragraphs):
        if i in blocks_dict:
            start_idx, end_idx, comp_name = blocks_dict[i]
            
            if comp_name in components:
                # Inserisci il codice aggiornato
                print(f"  ✓ Inserisco codice aggiornato per {comp_name}")
                
                for line in components[comp_name].split("\n"):
                    p = new_doc.add_paragraph(line)
                    # Formattazione codice
                    for run in p.runs:
                        run.font.name = "Consolas"
                        run.font.size = Pt(9)
                
                # Salta i vecchi paragrafi del blocco
                i = end_idx
                continue
            else:
                # Mantieni il blocco originale
                print(f"  ⚠️ Mantengo codice originale per {comp_name}")
        
        # Copia il paragrafo originale
        source_para = paragraphs[i]
        new_para = new_doc.add_paragraph(source_para.text)
        
        # Copia lo stile se possibile
        try:
            if source_para.style:
                new_para.style = source_para.style
        except:
            pass
        
        i += 1
    
    # Salva il documento
    print(f"\n💾 Salvataggio documento...")
    new_doc.save(OUTPUT_FILE)
    print(f"  ✓ Salvato: {OUTPUT_FILE}")
    
    # Genera anche il Markdown di riferimento
    print(f"\n📝 Generazione file Markdown di riferimento...")
    md_file = create_markdown_reference(components)
    print(f"  ✓ Salvato: {md_file}")
    
    return True


def create_markdown_reference(components: dict) -> Path:
    """Crea un file Markdown con il codice aggiornato per riferimento."""
    md_content = """# 📋 Codice Aggiornato - Love Sync

Questo file contiene il codice aggiornato per tutti i componenti.
Usa questo come riferimento per verificare l'aggiornamento del documento Word.

---

"""
    
    for comp_name, code in components.items():
        md_content += f"""
## 📄 {comp_name}.jsx

```jsx
{code}
```

---

"""
    
    output_md = OUTPUT_FILE.with_name("LoveSync_Codice_Aggiornato.md")
    output_md.write_text(md_content, encoding="utf-8")
    return output_md


if __name__ == "__main__":
    print("\n" + "🔄 " * 20)
    print("   SCRIPT AGGIORNAMENTO ARTICOLO LOVE SYNC")
    print("🔄 " * 20)
    
    success = update_word_document()
    
    if success:
        print(f"\n{'='*60}")
        print("✅ OPERAZIONE COMPLETATA CON SUCCESSO!")
        print(f"{'='*60}")
        print(f"\n📁 File generati:")
        print(f"  1. {OUTPUT_FILE}")
        print(f"  2. {OUTPUT_FILE.with_name('LoveSync_Codice_Aggiornato.md')}")
        print(f"\n💡 Suggerimento: Apri il file Word e verifica le modifiche")
    else:
        print(f"\n{'='*60}")
        print("❌ OPERAZIONE FALLITA")
        print(f"{'='*60}")
